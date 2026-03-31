"""DOCX document parser module."""

import logging
from pathlib import Path
from typing import Optional, Any
import docx
import mammoth  # type: ignore
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from docvert.models.document import (
    Document,
    Block,
    Heading,
    Image,
    Paragraph as DocvertParagraph,
    Table as DocvertTable,
)

logger = logging.getLogger(__name__)


class DocvertConfig:
    """A stub for the config class in case it doesn't exist yet."""

    pass


class DocxParser:
    """Parser for Microsoft Word (.docx) files.

    Extracts paragraphs, headings, and tables using python-docx with a fallback to mammoth.

    Args:
        config (Optional[Any]): Configuration object. Defaults to an empty DocvertConfig.
    """

    def __init__(self, config: Optional[Any] = None):
        """Initialize the DocxParser.

        Args:
            config: Optional configuration object.
        """
        self.config = config or DocvertConfig()

    def parse(self, file_path: str | Path) -> Document:
        """Parses a DOCX file into a Document object.

        Attempts to use python-docx first. If it fails, falls back to mammoth.

        Args:
            file_path (str | Path): Path to the DOCX file.

        Returns:
            Document: The parsed document structure.

        Raises:
            FileNotFoundError: If the specified file does not exist.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            return self._parse_with_python_docx(file_path)
        except Exception as e:
            logger.warning(f"python-docx failed: {e}. Falling back to mammoth.")
            return self._parse_with_mammoth(file_path)

    def _parse_with_python_docx(self, file_path: Path) -> Document:
        """Parses a DOCX file using python-docx.

        Args:
            file_path (Path): Path to the DOCX file.

        Returns:
            Document: The parsed document.
        """
        doc: _Document = docx.Document(str(file_path))
        docvert_doc = Document()

        for element in doc.element.body:
            if isinstance(element, CT_P):
                p = Paragraph(element, doc)
                blocks = self._process_paragraph(p)
                docvert_doc.blocks.extend(blocks)
            elif element.tag.endswith("tbl"):
                # Very basic table handling
                for table in doc.tables:
                    if table._element == element:
                        docvert_doc.blocks.append(self._process_table(table))
                        break

        return docvert_doc

    def _process_paragraph(self, p: Paragraph) -> list[Block]:
        """Processes a python-docx Paragraph into a list of Docvert Blocks.

        Determines if the paragraph is a heading or normal text based on styling and heuristics,
        and extracts any inline images.

        Args:
            p (Paragraph): The python-docx paragraph element.

        Returns:
            list[Block]: A list of corresponding Docvert blocks (Heading, Paragraph, and/or Image).
        """
        blocks: list[Block] = []
        text = p.text.strip()

        # Extract text block if any
        if text:
            score, level = self._calculate_heading_score(p)
            metadata = {"style": p.style.name if p.style else None, "score": score}

            if score >= 85:
                blocks.append(
                    Heading(content=text, level=level, score=score, metadata=metadata)
                )
            elif score >= 70:
                logger.warning(f"Borderline heading (score {score}): {text[:30]}...")
                blocks.append(DocvertParagraph(content=text, metadata=metadata))
            else:
                blocks.append(DocvertParagraph(content=text, metadata=metadata))

        # Extract images from runs
        for run in p.runs:
            drawings = run._element.findall(
                ".//*{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            if not drawings:
                # Fallback for mock tests without namespaces
                drawings = run._element.findall(".//drawing")

            for drawing in drawings:
                blips = drawing.findall(
                    ".//*{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                )
                if not blips:
                    # Fallback for mock tests
                    blips = drawing.findall(".//blip")

                for blip in blips:
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if not embed_id:
                        # Fallback for mock tests
                        embed_id = blip.get("embed")

                    if not embed_id and hasattr(blip, "attrib"):
                        # Sometimes lxml attrib is just 'embed'
                        embed_id = blip.attrib.get("embed") or blip.attrib.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )

                    # Get the part from the relationship ID
                    if (
                        embed_id
                        and hasattr(p, "part")
                        and hasattr(p.part, "related_parts")
                    ):
                        part = p.part.related_parts.get(embed_id)
                        if part:
                            # Map content type to extension
                            ext = self._get_extension_from_content_type(
                                part.content_type
                            )
                            blocks.append(
                                Image(
                                    content="",
                                    alt_text="extracted image",
                                    extension=ext,
                                    image_bytes=part.blob,
                                )
                            )

        return blocks

    def _get_extension_from_content_type(self, content_type: str) -> str:
        """Maps an image content type to a file extension.

        Args:
            content_type (str): The MIME type of the image.

        Returns:
            str: The corresponding file extension (e.g., '.png').
        """
        mapping = {
            "image/jpeg": ".jpg",
            "image/png": ".png",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/x-icon": ".ico",
            "image/svg+xml": ".svg",
            "image/webp": ".webp",
        }
        return mapping.get(content_type, ".bin")

    def _calculate_heading_score(self, p: Paragraph) -> tuple[int, int]:
        """Heuristics to determine if a paragraph is a heading.

        Args:
            p (Paragraph): The python-docx paragraph element.

        Returns:
            tuple[int, int]: A tuple containing the confidence score (0-100) and candidate heading level.
                score >= 85 -> heading
                score 70-84 -> warning, maybe heading
                score < 70 -> regular paragraph
        """
        score = 0
        level = 1

        style_name = p.style.name.lower() if p.style else ""

        # 1. Built-in styles
        if "heading" in style_name:
            score += 90
            try:
                # 'Heading 1' -> level 1
                level = int(style_name.split()[-1])
            except ValueError:
                level = 1

        # 2. Outline level (often 0 for body text, 1-9 for headings in word)
        # Using a simplified check, if outline_level is 0, it's body text (or top level depending on Word version)
        if p.style and hasattr(p.style, "paragraph_format"):
            outline_level = getattr(p.style.paragraph_format, "outline_level", None)
            # In docx, NO_OUTLINE_LEVEL is usually 9 (for some versions) or None. Let's do a basic check
            if outline_level is not None and outline_level < 9:
                score += 50
                level = outline_level + 1

        # 3. Formatting heuristics
        runs = p.runs
        if runs:
            # Check if majority of text is bold
            bold_chars = sum(len(r.text) for r in runs if r.bold)
            total_chars = sum(len(r.text) for r in runs)
            if total_chars > 0 and (bold_chars / total_chars) > 0.8:
                score += 30

            # Check font size
            sizes = [r.font.size.pt for r in runs if r.font and r.font.size]
            if sizes:
                avg_size = sum(sizes) / len(sizes)
                if avg_size >= 14:
                    score += 40
                elif avg_size >= 12:
                    score += 20

        # Max score is capped
        score = min(score, 100)

        return score, level

    def _process_table(self, table: Table) -> DocvertTable:
        """Processes a python-docx Table into a Docvert Table block.

        Args:
            table (Table): The python-docx table element.

        Returns:
            DocvertTable: The parsed table block.
        """
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_data.append(row_data)
        return DocvertTable(content="", rows=rows_data)

    def _parse_with_mammoth(self, file_path: Path) -> Document:
        """Fallback method using mammoth to parse DOCX to Markdown, then to a Document.

        Args:
            file_path (Path): Path to the DOCX file.

        Returns:
            Document: The parsed document structure.
        """
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            md_text = result.value

            # Simple conversion from markdown to Docvert Document
            docvert_doc = Document()
            for line in md_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("#"):
                    level = len(line.split(" ")[0])
                    content = line[level:].strip()
                    docvert_doc.blocks.append(
                        Heading(content=content, level=level, score=100)
                    )
                else:
                    docvert_doc.blocks.append(DocvertParagraph(content=line))

            return docvert_doc
